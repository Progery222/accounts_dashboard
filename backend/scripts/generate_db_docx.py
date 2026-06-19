"""Генерация docs/database-schema.docx — описание схемы БД dashboard."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "database-schema.docx"

# (поле, тип, что это, зачем)
TABLES: list[tuple[str, str, list[tuple[str, str, str, str]]]] = [
    (
        "profile",
        "Группа соцсетевых аккаунтов (логическая «папка» в UI).",
        [
            ("id", "BIGSERIAL PK", "Ключ профиля", "FK из account; неизменяемый идентификатор."),
            ("name", "VARCHAR(255)", "Название профиля", "Отображение в сайдбаре и фильтрах."),
            ("color", "VARCHAR(7)", "Цвет в UI", "HEX, по умолчанию #6366f1."),
            ("is_hidden", "BOOLEAN", "Скрыт ли профиль", "Скрытые профили и их аккаунты не показываются на главном экране (если не включён режим «показать скрытые»)."),
            ("created_at", "TIMESTAMPTZ", "Когда создан", "Аудит."),
            ("updated_at", "TIMESTAMPTZ", "Когда изменён", "Аудит."),
        ],
    ),
    (
        "owner",
        "Владелец аккаунта (второй уровень группировки, опционально).",
        [
            ("id", "BIGSERIAL PK", "Ключ владельца", "FK из account."),
            ("name", "VARCHAR(255)", "Имя владельца", "Фильтрация списка и scope автообновления."),
            ("color", "VARCHAR(7)", "Цвет в UI", "HEX для бейджа владельца."),
            ("created_at", "TIMESTAMPTZ", "Когда создан", "Аудит."),
            ("updated_at", "TIMESTAMPTZ", "Когда изменён", "Аудит."),
        ],
    ),
    (
        "account",
        "Отслеживаемый аккаунт соцсети.",
        [
            ("id", "BIGSERIAL PK", "Ключ аккаунта", "FK для постов, снимков, Apify jobs."),
            ("username", "VARCHAR(255)", "Ник / id на площадке", "Вместе с platform — уникальная пара."),
            ("platform", "VARCHAR(20)", "Площадка", "tiktok, instagram, youtube, telegram, x, threads, facebook, rumble, reddit."),
            ("profile_id", "BIGINT FK → profile", "Профиль", "SET NULL при удалении профиля."),
            ("owner_id", "BIGINT FK → owner", "Владелец", "SET NULL при удалении владельца; опционально."),
            ("display_name", "VARCHAR(255)", "Отображаемое имя", "С площадки, может отличаться от username."),
            ("avatar_url", "VARCHAR(1024)", "URL аватара на CDN", "Fallback, если локальный файл ещё не скачан."),
            ("avatar_file", "VARCHAR(512)", "Путь к локальному аватару", "Скачивается при refresh; upload_to accounts/avatars/%Y/%m/."),
            ("avatar_missing", "BOOLEAN", "Аватара нет на площадке", "Не пытаться скачивать повторно."),
            ("bio", "TEXT", "Описание профиля", "Из scrape."),
            ("follower_count", "BIGINT", "Подписчики", "Текущее значение; агрегируется в summary."),
            ("like_count", "BIGINT", "Лайки (сумма)", "Для TikTok/IG и др.; часть агрегатов."),
            ("view_count", "BIGINT", "Просмотры (сумма)", "Агрегируется из постов при refresh."),
            ("post_count", "INT", "Число постов", "Синхронизируется со списком постов."),
            ("link_click_count", "BIGINT", "Клики по ссылке в bio", "Из Links API; обновляется при refresh."),
            ("profile_unavailable", "BOOLEAN", "Профиль недоступен", "Удалён/заблокирован на площадке; фильтр refresh."),
            ("created_at", "TIMESTAMPTZ", "Когда добавлен", "Аудит; новые аккаунты без refresh выделяются в UI."),
            ("updated_at", "TIMESTAMPTZ", "Когда обновлён", "Меняется только при успешном refresh."),
            ("UNIQUE (username, platform)", "—", "Уникальность пары", "Один ник на площадке = одна строка; upsert при повторном добавлении."),
        ],
    ),
    (
        "account_snapshot",
        "Дневной снимок метрик аккаунта.",
        [
            ("id", "BIGSERIAL PK", "Ключ снимка", "—"),
            ("account_id", "BIGINT FK → account", "Аккаунт", "CASCADE при удалении аккаунта."),
            ("date", "DATE", "Календарная дата", "Один снимок на аккаунт в день."),
            ("follower_count", "BIGINT", "Подписчики на дату", "Для дельт 1/7/30 дней."),
            ("like_count", "BIGINT", "Лайки на дату", "Для дельт."),
            ("view_count", "BIGINT", "Просмотры на дату", "Для дельт и TV-графиков."),
            ("post_count", "INT", "Посты на дату", "Для дельт."),
            ("link_click_count", "BIGINT", "Клики на дату", "Для дельт."),
            ("UNIQUE (account_id, date)", "—", "Уникальность", "Идемпотентность take_snapshot_if_needed()."),
        ],
    ),
    (
        "post",
        "Пост отслеживаемого аккаунта.",
        [
            ("id", "BIGSERIAL PK", "Ключ поста", "FK для post_snapshot."),
            ("account_id", "BIGINT FK → account", "Аккаунт", "CASCADE."),
            ("external_id", "VARCHAR(255)", "ID поста на площадке", "Дедуп-ключ вместе с account_id."),
            ("description", "TEXT", "Текст / подпись", "Для аналитики хештегов."),
            ("hashtags", "JSONB", "Хештеги", "Массив строк; insights."),
            ("thumbnail_url", "VARCHAR(2048)", "URL превью", "CDN fallback."),
            ("thumbnail_file", "VARCHAR(512)", "Локальное превью", "upload_to posts/thumbnails/%Y/%m/."),
            ("thumbnail_missing", "BOOLEAN", "Превью нет", "Не скачивать повторно."),
            ("post_url", "VARCHAR(2048)", "Ссылка на пост", "Открытие в новой вкладке."),
            ("view_count", "BIGINT", "Просмотры", "Текущее значение."),
            ("like_count", "BIGINT", "Лайки", "—"),
            ("comment_count", "BIGINT", "Комментарии", "—"),
            ("share_count", "BIGINT", "Репосты", "—"),
            ("posted_at", "TIMESTAMPTZ", "Дата публикации", "Сортировка и аналитика по времени."),
            ("missing_from_scrape_at", "TIMESTAMPTZ", "Пропал из scrape", "Оранжевая лампочка; не удалять сразу."),
            ("created_at", "TIMESTAMPTZ", "Когда добавлен", "Аудит."),
            ("updated_at", "TIMESTAMPTZ", "Когда обновлён", "При sync постов."),
            ("UNIQUE (account_id, external_id)", "—", "Уникальность", "Один пост = одна строка."),
        ],
    ),
    (
        "post_snapshot",
        "Дневной снимок метрик поста.",
        [
            ("id", "BIGSERIAL PK", "Ключ снимка", "—"),
            ("post_id", "BIGINT FK → post", "Пост", "CASCADE."),
            ("date", "DATE", "Календарная дата", "Один снимок на пост в день."),
            ("view_count", "BIGINT", "Просмотры", "Дельты в analytics/top-posts."),
            ("like_count", "BIGINT", "Лайки", "—"),
            ("comment_count", "BIGINT", "Комментарии", "—"),
            ("UNIQUE (post_id, date)", "—", "Уникальность", "Идемпотентность take_snapshot_if_needed()."),
        ],
    ),
    (
        "refresh_schedule_config",
        "Singleton (pk=1). Настройки автообновления и расписания.",
        [
            ("id", "INT PK = 1", "Фиксированный ключ", "Одна строка на всё приложение."),
            ("enabled", "BOOLEAN", "Автообновление включено", "Регистрирует auto_refresh_* jobs в APScheduler."),
            ("mode", "VARCHAR(10)", "Режим расписания", "interval — каждые N часов; times — фиксированные часы."),
            ("interval_hours", "INT", "Интервал (часы)", "1–24, если mode=interval."),
            ("times", "JSONB", "Слоты времени", 'Например ["06:00", "12:00", "18:00", "00:00"].'),
            ("skip_recent_hours", "INT", "Пропуск недавно обновлённых", "Не трогать аккаунты, обновлённые за последние N часов."),
            ("refresh_warm_enabled", "BOOLEAN", "Прогрев Facebook", "Перед bulk/refresh_all/авто."),
            ("auto_refresh_csv_report", "BOOLEAN", "CSV-отчёт", "Сохранять после завершения автообновления."),
            ("auto_refresh_telegram_enabled", "BOOLEAN", "Telegram-отчёт", "Отправка в чаты после автообновления."),
            ("auto_refresh_telegram_chat_id", "VARCHAR(32)", "Один chat ID (устар.)", "Используйте auto_refresh_telegram_chat_ids."),
            ("auto_refresh_telegram_chat_ids", "JSONB", "Список chat ID", "Получатели Telegram-отчёта."),
            ("include_hidden_platform_accounts", "BOOLEAN", "Скрытые платформы", "Включать в автообновление."),
            ("include_hidden_profile_accounts", "BOOLEAN", "Скрытые профили", "Включать в автообновление."),
            ("include_unavailable_accounts", "BOOLEAN", "Недоступные аккаунты", "profile_unavailable=true."),
            ("auto_refresh_platforms", "JSONB", "Фильтр платформ", "Пусто — все; иначе список id."),
            ("auto_refresh_profile_ids", "JSONB", "Фильтр профилей", "Пусто — все; иначе id или «none»."),
            ("auto_refresh_owner_ids", "JSONB", "Фильтр владельцев", "Пусто — все; иначе id или «none»."),
            ("account_delta_period_days", "SMALLINT", "Период дельт в UI", "1, 7 или 30 календарных дней."),
        ],
    ),
    (
        "global_visibility_config",
        "Singleton (pk=1). Глобально скрытые платформы.",
        [
            ("id", "INT PK = 1", "Фиксированный ключ", "Одна строка."),
            ("hidden_platforms", "JSONB", "Список id платформ", "Не показываются в UI и summary (если не include_hidden)."),
        ],
    ),
    (
        "auto_refresh_state",
        "Singleton (pk=1). Прогресс текущего/последнего автообновления (bulk, scheduler, manual).",
        [
            ("id", "INT PK = 1", "Фиксированный ключ", "—"),
            ("is_running", "BOOLEAN", "Идёт прогон", "409 при конфликте с refresh_all."),
            ("source", "VARCHAR(32)", "Источник", "scheduler, bulk_refresh, manual."),
            ("cancel_requested", "BOOLEAN", "Запрошена остановка", "POST auto-refresh-stop."),
            ("total_accounts", "INT", "Всего в очереди", "Прогресс-бар."),
            ("processed_accounts", "INT", "Обработано", "—"),
            ("success_accounts", "INT", "Успешно", "—"),
            ("failed_accounts", "INT", "С ошибкой", "—"),
            ("current_account", "VARCHAR(255)", "Текущий аккаунт", "Отображение в UI."),
            ("started_at", "TIMESTAMPTZ", "Начало", "—"),
            ("finished_at", "TIMESTAMPTZ", "Конец", "—"),
            ("last_error", "TEXT", "Последняя ошибка", "—"),
            ("last_report_csv", "TEXT", "CSV-отчёт", "Текст для скачивания."),
            ("last_report_generated_at", "TIMESTAMPTZ", "Когда сгенерирован отчёт", "—"),
            ("last_telegram_error", "TEXT", "Ошибка Telegram", "—"),
            ("last_telegram_sent_at", "TIMESTAMPTZ", "Когда отправлен Telegram", "—"),
            ("last_auto_refresh_error_account_ids", "JSONB", "ID с ошибками", "Последний завершённый прогон."),
            ("run_detail", "JSONB", "Детали прогона", "worker_count, items по аккаунтам."),
            ("updated_at", "TIMESTAMPTZ", "Когда изменён", "—"),
        ],
    ),
    (
        "refresh_all_state",
        "Singleton (pk=1). Прогресс «Собрать всех» (POST /api/accounts/refresh_all/).",
        [
            ("id", "INT PK = 1", "Фиксированный ключ", "Отдельно от auto_refresh_state."),
            ("is_running", "BOOLEAN", "Идёт прогон", "—"),
            ("cancel_requested", "BOOLEAN", "Остановка", "—"),
            ("total_accounts", "INT", "Всего", "—"),
            ("processed_accounts", "INT", "Обработано", "—"),
            ("success_accounts", "INT", "Успешно", "—"),
            ("failed_accounts", "INT", "С ошибкой", "—"),
            ("current_account", "VARCHAR(255)", "Текущий", "—"),
            ("started_at", "TIMESTAMPTZ", "Начало", "—"),
            ("finished_at", "TIMESTAMPTZ", "Конец", "—"),
            ("last_error", "TEXT", "Ошибка", "—"),
            ("last_report_csv", "TEXT", "CSV-отчёт", "Отдельный отчёт refresh_all."),
            ("last_report_generated_at", "TIMESTAMPTZ", "Когда отчёт", "—"),
            ("run_detail", "JSONB", "Детали", "—"),
            ("updated_at", "TIMESTAMPTZ", "Когда изменён", "—"),
        ],
    ),
    (
        "scrape_backend_config",
        "Singleton (pk=1). Способ сбора данных по платформе.",
        [
            ("id", "INT PK = 1", "Фиксированный ключ", "—"),
            ("facebook_backend", "VARCHAR(16)", "Facebook", "playwright | apify."),
            ("tiktok_backend", "VARCHAR(16)", "TikTok", "playwright | apify."),
            ("instagram_backend", "VARCHAR(16)", "Instagram", "playwright | apify."),
            ("youtube_backend", "VARCHAR(16)", "YouTube", "playwright | apify."),
            ("reddit_backend", "VARCHAR(16)", "Reddit", "playwright | apify."),
            ("rumble_backend", "VARCHAR(16)", "Rumble", "playwright | apify."),
            ("facebook_fallback_enabled", "BOOLEAN", "FB fallback", "Playwright→Apify при rate limit."),
            ("tiktok_fallback_enabled", "BOOLEAN", "TT fallback", "Playwright→Apify при капче."),
            ("instagram_fallback_enabled", "BOOLEAN", "IG fallback", "—"),
            ("youtube_fallback_enabled", "BOOLEAN", "YT fallback", "—"),
            ("reddit_fallback_enabled", "BOOLEAN", "Reddit fallback", "—"),
            ("rumble_fallback_enabled", "BOOLEAN", "Rumble fallback", "—"),
            ("updated_at", "TIMESTAMPTZ", "Когда изменён", "—"),
        ],
    ),
    (
        "apify_refresh_job",
        "История асинхронного refresh через Apify.",
        [
            ("id", "BIGSERIAL PK", "Ключ задачи", "—"),
            ("account_id", "BIGINT FK → account", "Аккаунт", "CASCADE."),
            ("platform", "VARCHAR(32)", "Площадка", "Снимок на момент запуска."),
            ("username_snapshot", "VARCHAR(255)", "Ник на момент запуска", "Если username сменится."),
            ("status", "VARCHAR(16)", "Статус", "queued, starting, running, succeeded, failed, aborted."),
            ("apify_run_id", "VARCHAR(64)", "ID run в Apify", "Индекс; webhook completion."),
            ("apify_actor_id", "VARCHAR(255)", "Actor", "—"),
            ("apify_dataset_id", "VARCHAR(64)", "Dataset", "—"),
            ("apify_stages", "JSONB", "Этапы pipeline", "—"),
            ("trigger", "VARCHAR(32)", "Триггер", "manual, refresh_all, bulk, scheduler."),
            ("parent_batch_id", "UUID", "Пакет", "Группировка bulk/refresh_all."),
            ("started_at", "TIMESTAMPTZ", "Начало", "—"),
            ("finished_at", "TIMESTAMPTZ", "Конец", "—"),
            ("error_message", "TEXT", "Ошибка", "—"),
            ("run_detail_extra", "JSONB", "Доп. детали", "—"),
            ("normalized_preview", "JSONB", "Превью payload", "До _apply_refresh."),
            ("created_at", "TIMESTAMPTZ", "Создан", "—"),
            ("updated_at", "TIMESTAMPTZ", "Обновлён", "—"),
        ],
    ),
    (
        "auto_refresh_point",
        "Точки пульса просмотров для TV-графиков.",
        [
            ("id", "BIGSERIAL PK", "Ключ точки", "—"),
            ("measured_at", "TIMESTAMPTZ", "Время замера", "Индекс; каждые 30 мин (:00, :30)."),
            ("local_date", "DATE", "Календарная дата (МСК)", "Индекс."),
            ("source", "VARCHAR(32)", "Источник", "scheduler, interval pulse."),
            ("slot_label", "VARCHAR(32)", "Метка слота", "Например 14:30."),
            ("view_count_total", "BIGINT", "Сумма просмотров", "По всем аккаунтам."),
            ("view_delta_from_prev_point", "BIGINT", "Прирост с прошлой точки", "TV pulse."),
            ("view_delta_from_day_start", "BIGINT", "Прирост с начала дня", "TV pulse."),
            ("platform_deltas", "JSONB", "Дельты по платформам", "{'tiktok': 1200, ...}."),
        ],
    ),
]


def _set_cell_text(cell, text: str, *, bold: bool = False, size_pt: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)


def _add_table(doc: Document, rows: list[tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"
    headers = ("Поле", "Тип", "Что это", "Зачем")
    for i, text in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], text, bold=True)

    for row_idx, row_data in enumerate(rows, start=1):
        for col_idx, text in enumerate(row_data):
            _set_cell_text(table.rows[row_idx].cells[col_idx], text)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_heading("Dashboard — описание схемы БД", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run("PostgreSQL, Django app ").bold = False
    run = intro.add_run("accounts")
    run.bold = True
    intro.add_run(
        ". Префикс таблиц в БД: accounts_. "
        "Singleton-таблицы всегда содержат одну строку с id=1."
    )
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)

    # ER overview
    doc.add_heading("Связи (кратко)", level=1)
    er = doc.add_paragraph(
        "profile -- account -- post -- post_snapshot\n"
        "owner  /          |-- account_snapshot\n"
        "                  +-- apify_refresh_job"
    )
    er.paragraph_format.left_indent = Cm(0.5)
    for r in er.runs:
        r.font.name = "Courier New"
        r.font.size = Pt(9)

    for table_name, description, rows in TABLES:
        doc.add_heading(table_name, level=1)
        p = doc.add_paragraph(description)
        for r in p.runs:
            r.font.size = Pt(10)
            r.italic = True
        db_name = doc.add_paragraph()
        db_run = db_name.add_run(f"Таблица в PostgreSQL: accounts_{table_name}")
        db_run.font.size = Pt(9)
        db_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        _add_table(doc, rows)
        doc.add_paragraph()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    tmp = OUT.with_suffix(".tmp.docx")
    doc.save(tmp)
    tmp.replace(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Written: {path}")
